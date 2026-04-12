from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt

current_source_text = ""


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser()
    parser.add_argument("--source", required=True)
    parser.add_argument("--output-docx", required=True)
    return parser.parse_args()


def strip_md(text: str) -> str:
    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1 (\2)", text)
    text = text.replace("**", "")
    text = text.replace("`", "")
    return text.strip()


def add_markdown_runs(paragraph, text: str) -> None:
    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1 (\2)", text)
    parts = re.split(r"(\*\*.*?\*\*|`.*?`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            rpr = run._element.get_or_add_rPr()
            rfonts = rpr.rFonts
            rfonts.set(qn("w:ascii"), "Consolas")
            rfonts.set(qn("w:hAnsi"), "Consolas")
        else:
            paragraph.add_run(part)


def configure_styles(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(6)

    for style_name, size in [("Title", 18), ("Subtitle", 14)]:
        if style_name not in document.styles:
            document.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        style = document.styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.paragraph_format.space_after = Pt(6)

    for level, size in [(1, 14), (2, 13), (3, 12), (4, 12)]:
        style = document.styles[f"Heading {level}"]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)


def add_page_field(paragraph, field_name: str) -> None:
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_name

    fld_char_sep = OxmlElement("w:fldChar")
    fld_char_sep.set(qn("w:fldCharType"), "separate")

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_char_begin)
    run._r.append(instr)
    run._r.append(fld_char_sep)
    run._r.append(fld_char_end)


def extract_section(text: str, heading: str) -> str:
    pattern = rf"^## {re.escape(heading)}\s*$"
    match = re.search(pattern, text, flags=re.MULTILINE)
    if not match:
        return ""
    start = match.end()
    next_match = re.search(r"^## .+$", text[start:], flags=re.MULTILINE)
    end = start + next_match.start() if next_match else len(text)
    return text[start:end].strip()


def parse_table_block(lines: list[str]) -> list[list[str]]:
    rows: list[list[str]] = []
    for line in lines:
        line = line.strip()
        if not line.startswith("|"):
            continue
        parts = [strip_md(cell) for cell in line.strip("|").split("|")]
        rows.append(parts)
    if len(rows) >= 2 and all(set(cell) <= {"-", ":"} for cell in rows[1]):
        rows.pop(1)
    return rows


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def format_table(table) -> None:
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.15
        if row_index == 0:
            for cell in row.cells:
                shade_cell(cell, "D9EAF7")
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True


def add_table(document: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    table = document.add_table(rows=len(rows), cols=max(len(r) for r in rows))
    for r_idx, row in enumerate(rows):
        for c_idx, text in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.text = ""
            p = cell.paragraphs[0]
            add_markdown_runs(p, text)
    format_table(table)
    document.add_paragraph()


def add_cover_page(document: Document, text: str) -> None:
    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Powered by GPT - Software Quality Verification")
    run.bold = True

    subtitle = document.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Final Report for Software Testing Course")

    document.add_paragraph()

    info_table = document.add_table(rows=4, cols=2)
    info_table.style = "Table Grid"
    info_rows = [
        ("System Under Test", "Online Sales Management System"),
        ("Repository", "https://github.com/Alucard30Dec/Online-Sales-Management-System"),
        ("Appendix Link", "https://github.com/Alucard30Dec/Online-Sales-Management-System/tree/main/Report%20Test%20subject/SV00123-ATU-A01"),
        ("Submission Date", "April 11, 2026"),
    ]
    for idx, (label, value) in enumerate(info_rows):
        info_table.cell(idx, 0).text = label
        info_table.cell(idx, 1).text = value
    format_table(info_table)

    document.add_paragraph()
    team_heading = document.add_paragraph()
    team_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    team_heading.add_run("Team Members").bold = True

    team_table = document.add_table(rows=4, cols=3)
    team_rows = [
        ["Name", "Class", "Student ID"],
        ["Hoang Van Thien", "22D1ITE-SWE03", "225051915"],
        ["Nguyen Thanh Dat", "22D1ITE-SWE03", "225050896"],
        ["Le Quang Duy", "22D1ITE-SWE03", "225051169"],
    ]
    for r_idx, row in enumerate(team_rows):
        for c_idx, value in enumerate(row):
            team_table.cell(r_idx, c_idx).text = value
    format_table(team_table)
    document.add_page_break()


def add_record_of_changes(document: Document, text: str) -> None:
    document.add_heading("Record of Changes", level=1)
    section_text = extract_section(text, "Record Of Changes")
    rows = parse_table_block(section_text.splitlines())
    add_table(document, rows)


def add_toc_page(document: Document) -> None:
    document.add_heading("Table of Contents", level=1)
    body_start = current_source_text.find("# I. Overview")
    if body_start == -1:
        return
    headings = extract_toc_entries(current_source_text[body_start:])
    for level, text_value in headings:
        p = document.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.6 * (level - 1))
        add_markdown_runs(p, text_value)
    document.add_page_break()


def resolve_image(source_path: Path, line: str) -> tuple[str, Path] | None:
    match = re.match(r"!\[(.*?)\]\((.*?)\)", line.strip())
    if not match:
        return None
    alt_text, rel_path = match.groups()
    image_path = (source_path.parent / rel_path).resolve()
    return alt_text, image_path


def add_image(document: Document, source_path: Path, line: str) -> None:
    resolved = resolve_image(source_path, line)
    if not resolved:
        return
    alt_text, image_path = resolved
    if image_path.exists():
        pic = document.add_picture(str(image_path), width=Inches(6.0))
        document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption = document.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.add_run(alt_text).italic = True
        document.add_paragraph()
    else:
        p = document.add_paragraph()
        p.add_run(f"[Missing image] {alt_text}: {image_path}")


def add_list(document: Document, lines: list[str]) -> None:
    bullet_re = re.compile(r"^(\s*)[-*]\s+(.*)$")
    ordered_re = re.compile(r"^(\s*)(\d+)\.\s+(.*)$")
    for line in lines:
        bullet = bullet_re.match(line)
        ordered = ordered_re.match(line)
        if bullet:
            indent = len(bullet.group(1)) // 2
            p = document.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.63 * indent)
            p.paragraph_format.first_line_indent = Cm(0)
            p.add_run("• ")
            add_markdown_runs(p, bullet.group(2))
        elif ordered:
            indent = len(ordered.group(1)) // 2
            p = document.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.63 * indent)
            p.paragraph_format.first_line_indent = Cm(0)
            p.add_run(f"{ordered.group(2)}. ")
            add_markdown_runs(p, ordered.group(3))


def add_paragraph_text(document: Document, lines: list[str]) -> None:
    text = " ".join(line.strip() for line in lines if line.strip())
    if not text:
        return
    p = document.add_paragraph()
    add_markdown_runs(p, text)


def add_footer_page_numbers(document: Document) -> None:
    section = document.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Page ")
    add_page_field(p, "PAGE")


def extract_toc_entries(body_text: str) -> list[tuple[int, str]]:
    entries: list[tuple[int, str]] = []
    for line in body_text.splitlines():
        match = re.match(r"^(#{1,3})\s+(.*)$", line.strip())
        if not match:
            continue
        level = len(match.group(1))
        entries.append((level, strip_md(match.group(2))))
    return entries


def add_body(document: Document, source_path: Path, text: str) -> None:
    start = text.find("# I. Overview")
    if start == -1:
        raise RuntimeError("Could not find report body start marker '# I. Overview'.")
    lines = text[start:].splitlines()
    i = 0
    first_h1 = True
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        if not stripped:
            i += 1
            continue

        heading_match = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if heading_match:
            level = len(heading_match.group(1))
            text_value = strip_md(heading_match.group(2))
            if level == 1 and not first_h1:
                document.add_page_break()
            if level == 1:
                first_h1 = False
            document.add_heading(text_value, level=min(level, 4))
            i += 1
            continue

        if stripped.startswith("!["):
            add_image(document, source_path, stripped)
            i += 1
            continue

        if stripped.startswith("|"):
            table_lines: list[str] = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            add_table(document, parse_table_block(table_lines))
            continue

        if re.match(r"^\s*[-*]\s+", line) or re.match(r"^\s*\d+\.\s+", line):
            list_lines: list[str] = []
            while i < len(lines) and lines[i].strip() and (
                re.match(r"^\s*[-*]\s+", lines[i]) or re.match(r"^\s*\d+\.\s+", lines[i])
            ):
                list_lines.append(lines[i])
                i += 1
            add_list(document, list_lines)
            continue

        para_lines: list[str] = []
        while i < len(lines):
            candidate = lines[i]
            candidate_stripped = candidate.strip()
            if not candidate_stripped:
                break
            if re.match(r"^(#{1,6})\s+", candidate_stripped):
                break
            if candidate_stripped.startswith("![") or candidate_stripped.startswith("|"):
                break
            if re.match(r"^\s*[-*]\s+", candidate) or re.match(r"^\s*\d+\.\s+", candidate):
                break
            para_lines.append(candidate)
            i += 1
        add_paragraph_text(document, para_lines)


def build_report(source_path: Path, output_docx: Path) -> None:
    global current_source_text
    text = source_path.read_text(encoding="utf-8")
    current_source_text = text
    document = Document()
    configure_styles(document)
    add_footer_page_numbers(document)
    add_cover_page(document, text)
    add_record_of_changes(document, text)
    document.add_page_break()
    add_toc_page(document)
    add_body(document, source_path, text)
    output_docx.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_docx)


def main() -> None:
    args = parse_args()
    source_path = Path(args.source).resolve()
    output_docx = Path(args.output_docx).resolve()
    build_report(source_path, output_docx)


if __name__ == "__main__":
    main()
