from __future__ import annotations

import re
from pathlib import Path

from docx import Document


REPORT_PATHS = [
    Path(r"E:\Project\Online-Sales-Management-System\Report Test subject\Powered by GPT\MainReport\Powered by GPT - Software Quality Verification.docx"),
    Path(r"E:\Project\Online-Sales-Management-System\Report Test subject\Powered by GPT - A02\MainReport\Powered by GPT - Software Quality Verification.docx"),
    Path(r"E:\Project\Online-Sales-Management-System\Report Test subject\Powered by GPT - Software Quality Verification.docx"),
    Path(r"E:\Project\Online-Sales-Management-System\Report Test subject\SV00123-ATU-A01\MainReport\Powered by GPT - Software Quality Verification - Final Report.docx"),
]


EXACT_REPLACEMENTS = {
    "SOFTWARE QUALITY VERITIFICATION": "SOFTWARE QUALITY VERIFICATION",
    "Ho Chi Minh City, January 2026": "Ho Chi Minh City, April 2026",
    "LIST OF CONTENTS": "TABLE OF CONTENTS",
    "Performed by Student:": "Prepared by Students:",
    "OVERVIEW": "CHAPTER 1. OVERVIEW",
    "TEST PLAN": "CHAPTER 2. TEST PLAN",
    "TEST DESIGN & EXECUTION": "CHAPTER 3. TEST DESIGN & EXECUTION",
    "DEFECT REPORT & METRICS": "CHAPTER 4. DEFECT REPORT & METRICS",
    "CONCLUSION AND FUTURE WORK": "CHAPTER 5. CONCLUSION AND FUTURE WORK",
    "GitHub": "Appendix 1. GitHub",
    "Excel test cases": "Appendix 2. Excel test cases",
    "Test script and test data": "Appendix 3. Test script and test data",
    "Final results workbook and images": "Appendix 4. Final results workbook and images",
    "Automation video": "Appendix 5. Automation video",
    "Bug tracker screenshots": "Appendix 6. Bug tracker screenshots",
    "Crud validation for products, suppliers, and customers": "CRUD validation for products, suppliers, and customers",
    "Api status codes and payload validation": "API status codes and payload validation",
    "Unsupported api query values": "Unsupported API query values",
    "The target base url is reachable": "The target base URL is reachable",
    "Confirmed mismatches have been logged into the defect register and github Issues": "Confirmed mismatches have been logged into the defect register and GitHub Issues",
    "Open a github issue for reproducible confirmed defects only": "Open a GitHub Issue for reproducible confirmed defects only",
    "4 confirmed defects remain open at the reporting point, including 3 high-severity defects and 1 medium-severity defects.": "4 confirmed defects remain open at the reporting point, including 3 high-severity defects and 1 medium-severity defect.",
    "Bootstrap. (n.d.). Get started with Bootstrap. Bootstrap v5.3. Retrieved April 12, 2026, from https://getbootstrap.com/docs/5.3/getting-started/introduction/": "Bootstrap. (n.d.). Introduction. Bootstrap. https://getbootstrap.com/docs/5.3/getting-started/introduction/",
    "GitHub. (n.d.). About Git. GitHub Docs. Retrieved April 12, 2026, from https://docs.github.com/en/get-started/using-git/about-git": "GitHub. (n.d.). About Git. GitHub Docs. https://docs.github.com/en/get-started/using-git/about-git",
    "Microsoft. (2023, January 12). Migrations overview. Microsoft Learn. https://learn.microsoft.com/en-us/ef/core/managing-schemas/migrations/": "Microsoft. (2023, January 12). Migrations overview. Microsoft Learn. https://learn.microsoft.com/en-us/ef/core/managing-schemas/migrations/",
    "Microsoft. (2024, June 17). Overview of ASP.NET Core MVC. Microsoft Learn. https://learn.microsoft.com/en-us/aspnet/core/mvc/overview": "Microsoft. (2024, June 17). Overview of ASP.NET Core MVC. Microsoft Learn. https://learn.microsoft.com/en-us/aspnet/core/mvc/overview",
    "Microsoft. (2024, November 12). Overview of Entity Framework Core. Microsoft Learn. https://learn.microsoft.com/en-us/ef/core/": "Microsoft. (2024, November 12). Overview of Entity Framework Core. Microsoft Learn. https://learn.microsoft.com/en-us/ef/core/",
    "Microsoft. (2025, July 30). Overview of ASP.NET Core. Microsoft Learn. https://learn.microsoft.com/en-us/aspnet/core/overview": "Microsoft. (2025, July 30). Overview of ASP.NET Core. Microsoft Learn. https://learn.microsoft.com/en-us/aspnet/core/overview",
    "Microsoft. (2025, November 12). Introduction to Identity on ASP.NET Core. Microsoft Learn. https://learn.microsoft.com/en-us/aspnet/core/security/authentication/identity": "Microsoft. (2025, November 12). Introduction to Identity on ASP.NET Core. Microsoft Learn. https://learn.microsoft.com/en-us/aspnet/core/security/authentication/identity",
}


NUMBERED_HEADING_REPLACEMENTS = {
    "Project Information": "1.1. Project Information",
    "System overview": "1.1.1. System overview",
    "Main business modules": "1.1.2. Main business modules",
    "API overview": "1.1.3. API overview",
    "Project Team and Task Allocation": "1.2. Project Team and Task Allocation",
    "Team members": "1.2.1. Team members",
    "Responsibility matrix": "1.2.2. Responsibility matrix",
    "Test Scope": "2.1. Test Scope",
    "In-Scope": "2.1.1. In-Scope",
    "Out-of-Scope": "2.1.2. Out-of-Scope",
    "Test Strategy & Approach": "2.2. Test Strategy & Approach",
    "Manual testing approach": "2.2.1. Manual testing approach",
    "Basic automation approach": "2.2.2. Basic automation approach",
    "Black-box testing": "2.2.3. Black-box testing",
    "White-box testing": "2.2.4. White-box testing",
    "Entry and exit criteria": "2.2.5. Entry and exit criteria",
    "Test Environment": "2.3. Test Environment",
    "Hardware and software environment": "2.3.1. Hardware and software environment",
    "Browsers, OS, devices": "2.3.2. Browsers, OS, devices",
    "Test accounts and test data": "2.3.3. Test accounts and test data",
    "Tools used": "2.3.4. Tools used",
    "Test Scenarios": "3.1. Test Scenarios",
    "UI test scenarios": "3.1.1. UI test scenarios",
    "API test scenarios": "3.1.2. API test scenarios",
    "Edge cases and exception scenarios": "3.1.3. Edge cases and exception scenarios",
    "Test Case Specification": "3.2. Test Case Specification",
    "UI test cases": "3.2.1. UI test cases",
    "API test cases": "3.2.2. API test cases",
    "Test case allocation by member": "3.2.3. Test case allocation by member",
    "Execution evidence summary": "3.2.4. Execution evidence summary",
    "Bug Management Tool and Workflow": "4.1. Bug Management Tool and Workflow",
    "Selected tool": "4.1.1. Selected tool",
    "Issue workflow": "4.1.2. Issue workflow",
    "Severity and priority rules": "4.1.3. Severity and priority rules",
    "Defect Log": "4.2. Defect Log",
    "Critical and high severity defects": "4.2.1. Critical and high severity defects",
    "Medium and low severity defects": "4.2.2. Medium and low severity defects",
    "Screenshots and reproduction evidence": "4.2.3. Screenshots and reproduction evidence",
    "Test Summary Metrics": "4.3. Test Summary Metrics",
    "Total test cases executed": "4.3.1. Total test cases executed",
    "Pass/Fail/Blocked statistics": "4.3.2. Pass/Fail/Blocked statistics",
    "Defect count by severity": "4.3.3. Defect count by severity",
    "Defect ratio and observations": "4.3.4. Defect ratio and observations",
    "Achievements and compliance with requirements": "5.1. Achievements and compliance with requirements",
    "Challenges and limitations": "5.2. Challenges and limitations",
    "Suggestions for future enhancements": "5.3. Suggestions for future enhancements",
}


def extract_caption_entries(doc: Document, prefix: str) -> list[str]:
    entries: list[str] = []
    in_block = False
    for para in doc.paragraphs:
        text = para.text.strip()
        style = para.style.name if para.style else ""
        if style == "Heading 1" and text == f"LIST OF {prefix.upper()}S":
            in_block = True
            continue
        if in_block and style == "Heading 1":
            break
        if in_block and style == "table of figures" and text.startswith(prefix):
            entries.append(text.split("\t")[0].strip())
    return entries


def patch_caption_texts(doc: Document, prefix: str, entries: list[str]) -> None:
    idx = 0
    for para in doc.paragraphs:
        text = para.text.strip()
        style = para.style.name if para.style else ""
        if style == "Caption" and text.startswith(f"{prefix} -:"):
            if idx < len(entries):
                para.text = entries[idx]
                idx += 1


def apply_exact_replacements(doc: Document) -> None:
    for para in doc.paragraphs:
        text = para.text.strip()
        if text in EXACT_REPLACEMENTS:
            para.text = EXACT_REPLACEMENTS[text]
        elif text in NUMBERED_HEADING_REPLACEMENTS and para.style.name in {"Heading 2", "Heading 3"}:
            para.text = NUMBERED_HEADING_REPLACEMENTS[text]


def clear_unneeded_paragraphs(doc: Document) -> None:
    for para in doc.paragraphs:
        text = para.text.strip()
        if text.startswith("Figma design link: No verified Figma URL was found in the current repository artifacts."):
            para.text = ""


def count_caption_placeholders(doc: Document) -> int:
    count = 0
    for para in doc.paragraphs:
        text = para.text.strip()
        if text.startswith("Table -:") or text.startswith("Figure -:"):
            count += 1
    return count


def main() -> None:
    for path in REPORT_PATHS:
        if not path.exists():
            continue

        doc = Document(str(path))

        table_entries = extract_caption_entries(doc, "Table")
        figure_entries = extract_caption_entries(doc, "Figure")

        apply_exact_replacements(doc)
        clear_unneeded_paragraphs(doc)
        patch_caption_texts(doc, "Table", table_entries)
        patch_caption_texts(doc, "Figure", figure_entries)

        if count_caption_placeholders(doc) != 0:
            raise RuntimeError(f"Unpatched caption placeholders remain in {path}")

        doc.save(str(path))
        print(path)
        print(path.stat().st_size)


if __name__ == "__main__":
    main()
