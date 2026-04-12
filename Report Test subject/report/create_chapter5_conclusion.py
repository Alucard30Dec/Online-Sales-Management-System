from __future__ import annotations

import csv
import os
from collections import Counter
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt


def shade_cell(cell, fill: str = "D9EAF7") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def configure_document(doc: Document) -> None:
    sec = doc.sections[0]
    sec.top_margin = Cm(2.54)
    sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(2.54)
    sec.right_margin = Cm(2.54)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(6)

    for level, size in [(1, 16), (2, 14), (3, 13)]:
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = True
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(4)


def add_paragraph(doc: Document, text: str, align=None) -> None:
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="Normal")
        p.paragraph_format.left_indent = Cm(0.63)
        p.paragraph_format.first_line_indent = Cm(-0.3)
        run = p.add_run("- " + item)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def add_table(
    doc: Document,
    title: str,
    headers: list[str],
    rows: list[list[str]],
    col_widths: list[float] | None = None,
) -> None:
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(title)
    r.italic = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(11)
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        shade_cell(hdr[i])
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)

    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)

    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(11)
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.15
    doc.add_paragraph()


def add_figure(doc: Document, img_path: Path, caption: str, width_inches: float = 5.9) -> None:
    if not img_path.exists():
        return
    doc.add_picture(str(img_path), width=Inches(width_inches))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    r.italic = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(11)
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    doc.add_paragraph()


def read_csv(path: Path) -> list[dict[str, str]]:
    with path.open(encoding="utf-8-sig", newline="") as f:
        return list(csv.DictReader(f))


def top_level_items(path: Path) -> list[str]:
    return sorted(p.name for p in path.iterdir())


def main() -> None:
    out_dir = Path(os.environ["TARGET_DIR"])
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / "Chapter 5 - Conclusion and Future Work.docx"

    clean_root = Path(os.environ["CLEAN_ROOT"])
    ui_csv = Path(os.environ["UI_CSV"])
    api_csv = Path(os.environ["API_CSV"])
    metrics_csv = Path(os.environ["METRICS_CSV"])
    defect_csv = Path(os.environ["DEFECT_CSV"])
    img_metrics = Path(os.environ["IMG_METRICS"])
    img_issue = Path(os.environ["IMG_ISSUE"])
    img_api = Path(os.environ["IMG_API"])

    ui_rows = read_csv(ui_csv)
    api_rows = read_csv(api_csv)
    metrics_rows = read_csv(metrics_csv)
    defect_rows = read_csv(defect_csv)

    metrics = {row["Metric"]: row["Value"] for row in metrics_rows}
    owner_counts = Counter(r["Owner"] for r in ui_rows)
    confirmed = [r for r in defect_rows if r["Record Type"] == "Confirmed Defect"]
    observations = [r for r in defect_rows if r["Record Type"] == "Observation"]
    open_confirmed = [r for r in confirmed if r["Current Status"] == "Open - Confirmed"]
    severity_counts = Counter(r["Severity"] for r in confirmed)

    clean_items = top_level_items(clean_root)

    total_cases = int(metrics["Total Test Cases"])
    executed_cases = int(metrics["Executed Test Cases"])
    passed_cases = int(metrics["Pass"])
    failed_cases = int(metrics["Fail"])
    blocked_cases = int(metrics["Blocked"])
    not_run_cases = int(metrics["Not Run"])
    scenario_count = int(metrics["Scenario Count In Phase 3"])
    mapped_scenarios = int(metrics["Scenarios Mapped To Test Cases"])
    executed_scenarios = int(metrics["Executed Scenarios"])
    pass_rate = metrics["Pass Rate On Executed %"]
    fail_rate = metrics["Fail Rate On Executed %"]
    execution_progress = metrics["Execution Progress %"]
    scenario_coverage = metrics["Scenario Execution Coverage %"]
    confirmed_count = len(confirmed)
    observation_count = len(observations)
    defect_ratio = round((confirmed_count / total_cases) * 100, 2) if total_cases else 0

    doc = Document()
    configure_document(doc)

    h = doc.add_paragraph(style="Heading 1")
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h.add_run("CHAPTER 5. CONCLUSION AND FUTURE WORK").bold = True

    h = doc.add_paragraph(style="Heading 2")
    h.add_run("5.1. Achievements and compliance with requirements").bold = True
    add_paragraph(
        doc,
        "The final submission achieved the main academic and practical goals of the Software Testing project. "
        "The team produced a full testing package that includes a structured report, presentation slides, scenario and test-case workbooks, final execution results, metrics, "
        "defect records, automation assets, supporting screenshots, runner outputs, and an automation demo video. These artifacts were synchronized so that major claims in the report "
        "can be traced back to an execution file, an evidence screenshot, or a live defect record.",
    )
    add_paragraph(
        doc,
        f"From an execution perspective, the project reached a complete recorded baseline with {executed_cases}/{total_cases} executed test cases, "
        f"{passed_cases} passing cases, {failed_cases} failing cases, {blocked_cases} blocked cases, and {not_run_cases} cases left as Not Run. "
        f"Scenario coverage also reached {executed_scenarios}/{scenario_count}, which means the documented scenario set was fully mapped and executed in the current evidence baseline.",
    )
    add_paragraph(
        doc,
        "The team also met the mandatory course requirements regarding individual contribution and defect handling. UI test-case ownership remained non-overlapping across members, "
        "each member contributed more than ten UI test cases, and confirmed defects were managed through a real GitHub Issues workflow with severity, priority, screenshots, and reproducible steps.",
    )

    add_table(
        doc,
        "Table 5.1. Final requirement compliance summary",
        ["Requirement area", "Current status", "Supporting evidence"],
        [
            ["Group-based final testing project", "Satisfied", "Team allocation tables and shared submission package"],
            ["At least 10 non-overlapping test cases per member", "Satisfied", f"UI ownership counts: {', '.join(f'{k}={v}' for k, v in owner_counts.items())}"],
            ["Industry-style test case fields", "Satisfied", "UI and API Excel/CSV test-case files"],
            ["Execution evidence", "Satisfied", "Screenshots, TRX, TXT, XML, and final result workbook"],
            ["Real bug-management tool", "Satisfied", "4 live GitHub Issues linked from the defect register"],
            ["Automation bonus support", "Satisfied", "Selenium UI runs, Newman outputs, and MP4 demo video"],
            ["Submission packaging", "Satisfied", ", ".join(clean_items)],
        ],
        [5.0, 3.4, 8.6],
    )

    add_table(
        doc,
        "Table 5.2. Final delivery baseline",
        ["Measure", "Value", "Interpretation"],
        [
            ["Total test cases", str(total_cases), "Combined UI and API baseline"],
            ["Executed test cases", str(executed_cases), f"Execution progress {execution_progress}%"],
            ["Pass / Fail / Blocked", f"{passed_cases} / {failed_cases} / {blocked_cases}", f"Pass rate {pass_rate}%, fail rate {fail_rate}%"],
            ["Scenario mapping", f"{mapped_scenarios}/{scenario_count}", "All documented scenarios mapped to test cases"],
            ["Executed scenarios", f"{executed_scenarios}/{scenario_count}", f"Scenario execution coverage {scenario_coverage}%"],
            ["Confirmed defects", str(confirmed_count), "Backed by GitHub Issues and screenshots"],
        ],
        [4.3, 3.0, 9.7],
    )
    add_figure(
        doc,
        img_metrics,
        "Figure 5.1. Final execution and metrics summary used to support the conclusion of the testing project.",
    )

    h = doc.add_paragraph(style="Heading 2")
    h.add_run("5.2. Challenges and limitations").bold = True
    add_paragraph(
        doc,
        "Although the submission achieved full execution coverage for the documented test-case baseline, several limitations remain. "
        "These limitations do not invalidate the test results, but they explain why the current package should be described as a strong and evidence-backed submission rather than as proof that the product is defect-free.",
    )
    add_bullets(
        doc,
        [
            f"{confirmed_count} confirmed defects remain open at the reporting point, including {severity_counts.get('High', 0)} high-severity defects and {severity_counts.get('Medium', 0)} medium-severity defects.",
            "The project captured defect discovery and live issue tracking, but post-fix retest evidence is still unavailable because the underlying product defects were not fixed within the current cycle.",
            "Cross-browser evidence exists only at a basic smoke level on Microsoft Edge; it is not a full compatibility certification matrix.",
            "The environment is a controlled local execution baseline rather than a production-like staging environment.",
            "Non-functional areas such as load, penetration, accessibility, and recovery testing were intentionally excluded from scope.",
        ],
    )
    add_paragraph(
        doc,
        "The most important practical limitation is that the defect lifecycle is not fully closed. The team successfully identified, documented, and classified defects, but closure-level verification still depends on future product fixes followed by a controlled retest cycle.",
    )

    add_table(
        doc,
        "Table 5.3. Current limitations and impact",
        ["Limitation", "Current impact", "Reason it remains open"],
        [
            ["Open confirmed defects", "Product cannot be claimed defect-free", "Fixes were not completed before the final reporting point"],
            ["No post-fix pass retest", "Defect lifecycle is not fully closed", "No new build with verified fixes was available"],
            ["Limited browser matrix", "Compatibility claims must remain conservative", "Only Chrome baseline plus Edge smoke were executed"],
            ["Local environment only", "Deployment-specific behavior was not certified", "The course project was executed in a local environment"],
            ["Out-of-scope non-functional testing", "Performance and security conclusions are limited", "These areas were outside the official functional scope"],
        ],
        [4.8, 5.8, 6.4],
    )
    add_figure(
        doc,
        img_issue,
        "Figure 5.2. Example of a confirmed open defect that remains unresolved at the final reporting point.",
    )

    h = doc.add_paragraph(style="Heading 2")
    h.add_run("5.3. Suggestions for future enhancements").bold = True
    add_paragraph(
        doc,
        "Future work should focus first on product risk reduction and then on process maturity. The next cycle should prioritize defect fixing and closure, because this is the shortest path to improving actual software quality rather than only improving documentation quality.",
    )
    add_bullets(
        doc,
        [
            "Fix the four confirmed open defects and perform controlled post-fix retests with updated evidence.",
            "Expand browser coverage beyond Chrome and Edge smoke to include a more defendable compatibility baseline.",
            "Add more negative and resilience-oriented API tests, especially for invalid filters, unsupported values, and boundary conditions.",
            "Broaden automation coverage for high-value regression flows while keeping outputs mapped to the official test-case IDs.",
            "Introduce CI-based automation execution so that UI and API smoke runs can be reproduced more consistently after future changes.",
        ],
    )
    add_paragraph(
        doc,
        "From an academic perspective, the submission already demonstrates structured planning, evidence-based execution, real defect management, and practical automation support. "
        "From an engineering perspective, the next iteration should move from defect discovery maturity to defect-closure maturity.",
    )

    add_table(
        doc,
        "Table 5.4. Recommended future enhancement roadmap",
        ["Priority", "Enhancement", "Expected benefit"],
        [
            ["1", "Fix open confirmed defects and rerun retests", "Closes the defect lifecycle and improves product stability"],
            ["2", "Extend browser compatibility verification", "Improves confidence in UI consistency across browsers"],
            ["3", "Increase API negative-path automation", "Strengthens backend validation coverage"],
            ["4", "Expand UI regression automation", "Improves repeatability for critical business flows"],
            ["5", "Move automation to CI or scheduled reruns", "Improves repeatable verification after changes"],
        ],
        [2.0, 7.5, 7.0],
    )
    add_figure(
        doc,
        img_api,
        "Figure 5.3. Existing automation evidence that can be expanded in future iterations to increase regression strength.",
    )

    add_paragraph(
        doc,
        "Overall, the final submission meets the required deliverables and demonstrates a credible software-testing process with strong traceability between scenarios, test cases, execution evidence, metrics, and GitHub-based defect records. "
        "The package is submission-ready and academically defensible. The remaining improvement opportunities are primarily product-fix and regression-depth tasks, not missing-report or missing-artifact problems.",
    )

    doc.save(out_path)
    print(out_path)
    print(out_path.stat().st_size)


if __name__ == "__main__":
    main()
