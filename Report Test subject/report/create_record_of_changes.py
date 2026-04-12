from __future__ import annotations

import os
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


def shade_cell(cell, fill: str = "D9EAF7") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def configure_document(doc: Document) -> None:
    sec = doc.sections[0]
    sec.top_margin = Cm(2.54)
    sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(2.54)
    sec.right_margin = Cm(2.54)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(6)

    heading = doc.styles["Heading 1"]
    heading.font.name = "Times New Roman"
    heading.font.size = Pt(16)
    heading.font.bold = True
    heading._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    heading.paragraph_format.space_before = Pt(10)
    heading.paragraph_format.space_after = Pt(4)


def add_paragraph(doc: Document, text: str, align=None) -> None:
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def add_table(doc: Document, title: str, headers: list[str], rows: list[list[str]], col_widths: list[float]) -> None:
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(title)
    r.italic = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(11)
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        shade_cell(hdr[i])

    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value

    for row in table.rows:
        for i, width in enumerate(col_widths):
            row.cells[i].width = Cm(width)
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(11)
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.15


def main() -> None:
    target_dir = Path(os.environ["TARGET_DIR"])
    target_dir.mkdir(parents=True, exist_ok=True)
    out_path = target_dir / "Record of Changes.docx"

    doc = Document()
    configure_document(doc)

    heading = doc.add_paragraph(style="Heading 1")
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.add_run("RECORD OF CHANGES").bold = True

    add_paragraph(
        doc,
        "This section summarizes the major revision milestones of the final report and the submission package. "
        "Only meaningful changes that affected scope, execution, evidence, defects, metrics, or final packaging are recorded here.",
    )

    add_table(
        doc,
        "Table R.1. Record of changes",
        ["Version", "Date", "Author", "Change Summary"],
        [
            ["0.1", "2026-04-05", "QA Team", "Established the report structure, completed the initial project audit, and defined the first testing scope and plan baseline."],
            ["0.5", "2026-04-05", "QA Team", "Added the UI and API test cases, prepared the test data set, and documented the initial automation approach."],
            ["0.8", "2026-04-06", "QA Team", "Added execution evidence, initial defect records, early metrics, and supporting automation outputs."],
            ["1.0", "2026-04-06", "Hoang Van Thien", "Consolidated the main report content, appendix linkage, and final submission package structure."],
            ["1.1", "2026-04-10", "QA Team", "Expanded verified UI execution coverage and synchronized the report with the updated results, metrics, and evidence set."],
            ["1.2", "2026-04-11", "QA Team", "Removed the remaining Not Run status, synchronized the live GitHub Issues, and refreshed the defect and metrics sections."],
            ["1.3", "2026-04-12", "QA Team", "Finalized chapter wording, references, appendix content, formatting consistency, and submission-ready packaging."],
        ],
        [2.0, 3.0, 3.4, 10.6],
    )

    add_paragraph(
        doc,
        "In the final report, this section should appear immediately after the cover page and before the table of contents.",
    )

    doc.save(out_path)
    print(out_path)
    print(out_path.stat().st_size)


if __name__ == "__main__":
    main()
