from __future__ import annotations

import csv
import os
from collections import Counter
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt


def shade_cell(cell, fill: str = "D9EAF7") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def configure_document(doc: Document) -> None:
    sec = doc.sections[0]
    sec.top_margin = Cm(2.54)
    sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(2.54)
    sec.right_margin = Cm(2.54)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(6)

    for level, size in [(1, 16), (2, 14), (3, 13)]:
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = True
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(4)


def add_paragraph(doc: Document, text: str, align=None) -> None:
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="Normal")
        p.paragraph_format.left_indent = Cm(0.63)
        p.paragraph_format.first_line_indent = Cm(-0.3)
        run = p.add_run("• " + item)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def add_table(
    doc: Document,
    title: str,
    headers: list[str],
    rows: list[list[str]],
    col_widths: list[float] | None = None,
) -> None:
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(title)
    r.italic = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(11)
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        shade_cell(hdr[i])
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)

    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)

    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(11)
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.15
    doc.add_paragraph()


def add_figure(doc: Document, img_path: Path, caption: str, width_inches: float = 5.8) -> None:
    if not img_path.exists():
        return
    doc.add_picture(str(img_path), width=Inches(width_inches))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    r.italic = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(11)
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    doc.add_paragraph()


def read_csv(path: Path) -> list[dict[str, str]]:
    with path.open(encoding="utf-8-sig", newline="") as f:
        return list(csv.DictReader(f))


def scenario_rows_from_md(path: Path) -> list[list[str]]:
    text = path.read_text(encoding="utf-8")
    rows: list[list[str]] = []
    for line in text.splitlines():
        if line.startswith("| `SCN-"):
            rows.append([c.strip().strip("`") for c in line.strip("|").split("|")])
    return rows


def get_case(rows: list[dict[str, str]], tcid: str) -> dict[str, str]:
    return next(r for r in rows if r.get("Test Case ID") == tcid)


def sample_case_rows(case: dict[str, str], fields: list[str]) -> list[list[str]]:
    return [[field, case.get(field, "")] for field in fields]


def main() -> None:
    out_dir = Path(os.environ["TARGET_DIR"])
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / "Chapter 3 - Test Design and Execution.docx"

    ui_csv = Path(os.environ["UI_CSV"])
    api_csv = Path(os.environ["API_CSV"])
    final_csv = Path(os.environ["FINAL_CSV"])
    scenarios_md = Path(os.environ["SCENARIOS_MD"])
    img_login = Path(os.environ["IMG_LOGIN"])
    img_invoice = Path(os.environ["IMG_INVOICE_FAIL"])
    img_import = Path(os.environ["IMG_IMPORT"])
    img_api = Path(os.environ["IMG_API"])

    ui_rows = read_csv(ui_csv)
    api_rows = read_csv(api_csv)
    final_rows = read_csv(final_csv)
    scenario_rows = scenario_rows_from_md(scenarios_md)

    ui_case_pass = get_case(ui_rows, "TC-UI-AUTH-001")
    ui_case_fail = get_case(ui_rows, "TC-UI-INV-001")
    api_case = get_case(api_rows, "TC-API-HLT-001")

    scenario_interface_counts = Counter(r[2] for r in scenario_rows)
    ui_modules = Counter(r["Module"] for r in ui_rows)
    api_modules = Counter(r["Module"] for r in api_rows)
    owner_counts = Counter(r["Owner"] for r in ui_rows)
    final_exec_counts = Counter(r["Execution Status"] for r in final_rows)
    final_interface_counts = Counter(r["Interface"] for r in final_rows)

    doc = Document()
    configure_document(doc)

    h = doc.add_paragraph(style="Heading 1")
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h.add_run("CHAPTER 3. TEST DESIGN & EXECUTION").bold = True

    h = doc.add_paragraph(style="Heading 2")
    h.add_run("3.1. Test Scenarios").bold = True
    add_paragraph(
        doc,
        "This chapter presents how the test scenarios and test cases were designed for the Online Sales Management System and how the resulting execution evidence supports the final submission. "
        "The design phase was based on real system modules, route behavior, role boundaries, validation rules, and business-risk areas identified through source inspection and local execution.",
    )
    add_paragraph(
        doc,
        "The scenario set was intentionally broader than the minimum number of test cases required by the course. This allowed the team to split ownership cleanly across members, "
        "cover both positive and negative flows, and preserve enough business-rule depth to support defect discovery and traceability.",
    )

    h = doc.add_paragraph(style="Heading 3")
    h.add_run("3.1.1. UI test scenarios").bold = True
    add_paragraph(
        doc,
        "The UI scenario design covers the administrative interface and the public product catalog. Administrative scenarios dominate the scope because the highest-risk workflows of the system "
        "exist in secured modules such as authentication, permissions, purchases, invoices, stock, reports, and product import. Public UI scenarios were included to verify search, sorting, filtering, and product-detail behavior.",
    )
    add_paragraph(
        doc,
        f"Based on the current scenario matrix, the project contains {scenario_interface_counts['Admin UI']} Admin UI scenarios and {scenario_interface_counts['Public UI']} Public UI scenarios. "
        "These scenarios were then split into detailed executable UI test cases with non-overlapping ownership.",
    )
    add_table(
        doc,
        "Table 3.1. UI scenario distribution by interface",
        ["Interface", "Scenario count", "Main focus"],
        [
            ["Admin UI", str(scenario_interface_counts["Admin UI"]), "Authentication, permissions, CRUD, purchases, invoices, stock, reports, import"],
            ["Public UI", str(scenario_interface_counts["Public UI"]), "Search, filtering, sorting, and product details"],
            ["Total UI scenarios", str(scenario_interface_counts["Admin UI"] + scenario_interface_counts["Public UI"]), "Complete UI scenario coverage"],
        ],
        [3.2, 3.2, 10.6],
    )
    add_table(
        doc,
        "Table 3.2. Representative UI scenario groups",
        ["Module", "Representative scenario focus", "Risk level"],
        [
            ["Authentication", "Valid login, invalid login, inactive-account blocking", "Critical"],
            ["Permissions", "Role-based denial and module visibility", "Critical"],
            ["Products / Import", "Validation, duplicate handling, preview and confirm behavior", "Critical"],
            ["Purchases", "Draft creation, validation, receiving, cancellation rules", "Critical"],
            ["Invoices", "Creation, payment transition, cancellation, stock-sensitive logic", "Critical"],
            ["Stock / Reports", "Visibility, filtering, export-source verification", "High"],
            ["Public Catalog", "Search, filter, sort, and product details", "Medium"],
        ],
        [4.0, 10.2, 2.8],
    )
    add_figure(
        doc,
        img_login,
        "Figure 3.1. Example UI execution evidence for a successful authentication scenario.",
    )

    h = doc.add_paragraph(style="Heading 3")
    h.add_run("3.1.2. API test scenarios").bold = True
    add_paragraph(
        doc,
        "API scenario design focused on the real endpoints exposed by the project rather than hypothetical or undocumented services. The scenario set covers the Health API and the public Catalog API. "
        "API scenarios were selected to verify service availability, response structure, query handling, detail retrieval, and invalid-parameter behavior.",
    )
    add_paragraph(
        doc,
        f"The current scenario baseline contains {scenario_interface_counts['API']} API scenarios, which were expanded into {len(api_rows)} detailed API test cases in the dedicated workbook. "
        "This split allowed the team to test more than one parameter combination under a single high-level scenario when needed.",
    )
    add_table(
        doc,
        "Table 3.3. API scenario coverage",
        ["API scenario area", "Scenario focus", "Representative endpoint"],
        [
            ["Health API", "Service availability and payload structure", "GET /api/v1/health"],
            ["Catalog list", "Pagination, search, filter, and sort behavior", "GET /api/v1/catalog/products"],
            ["Catalog detail", "Valid ID and not-found behavior", "GET /api/v1/catalog/products/{id}"],
            ["Catalog metadata", "Trending and filter-data retrieval", "GET /api/v1/catalog/trending and /filters"],
            ["Validation handling", "Unsupported or invalid query combinations", "Catalog endpoints with invalid parameters"],
        ],
        [4.0, 9.0, 4.0],
    )
    add_figure(
        doc,
        img_api,
        "Figure 3.2. API scenario execution evidence produced by the Newman full-run output.",
    )

    h = doc.add_paragraph(style="Heading 3")
    h.add_run("3.1.3. Edge cases and exception scenarios").bold = True
    add_paragraph(
        doc,
        "Edge cases and exception scenarios were intentionally included to prevent the design from being limited to happy-path verification. "
        "These scenarios are important because the rubric rewards logical depth, error handling, and the ability to identify hidden business-rule defects.",
    )
    add_bullets(
        doc,
        [
            "inactive-account login despite otherwise valid credentials",
            "role-based access attempts to secured modules",
            "duplicate SKU and invalid numeric values in product management",
            "non-xlsx and malformed uploads for product import",
            "expired or missing preview state during import confirmation",
            "purchase creation with missing supplier or empty item list",
            "invoice creation with insufficient stock or invalid posted state",
            "invoice cancellation and payment transitions under sensitive business states",
            "unsupported sort and invalid query combinations in API requests",
        ],
    )
    add_table(
        doc,
        "Table 3.4. Representative edge and exception scenarios",
        ["Scenario ID", "Area", "Exception / edge focus", "Why it matters"],
        [
            ["SCN-AUTH-003", "Authentication", "Inactive account login attempt", "Prevents unauthorized account reuse"],
            ["SCN-GOV-001", "Permissions", "Sales role opens restricted module", "Verifies authorization boundaries"],
            ["SCN-PROD-006", "Product Import", "Invalid upload type or malformed workbook", "Protects the system from invalid bulk data input"],
            ["SCN-PROD-007", "Product Import", "Preview-confirm lifecycle issues", "Targets high-risk cache and state behavior"],
            ["SCN-PUR-002", "Purchases", "Missing supplier or missing item list", "Protects transaction integrity"],
            ["SCN-INV-003", "Invoices", "Server-side price and transaction safety", "Targets hidden business-rule and security risk"],
            ["SCN-INV-006", "Invoices", "Cancellation and stock-return logic", "Directly affects inventory consistency"],
            ["SCN-API-003", "Catalog API", "Invalid sort / page / page-size validation", "Protects API robustness and predictable error handling"],
        ],
        [3.0, 3.0, 5.4, 6.6],
    )
    add_figure(
        doc,
        img_import,
        "Figure 3.3. Example edge-case evidence from the product import preview validation flow.",
    )

    h = doc.add_paragraph(style="Heading 2")
    h.add_run("3.2. Test Case Specification").bold = True
    add_paragraph(
        doc,
        "After the scenario set was stabilized, each scenario was expanded into detailed executable test cases. The final design uses separate UI and API workbooks, with consistent fields for identifiers, "
        "preconditions, test data, steps, expected results, actual results, ownership, and evidence references. This structure supports both execution discipline and post-execution traceability.",
    )

    h = doc.add_paragraph(style="Heading 3")
    h.add_run("3.2.1. UI test cases").bold = True
    add_paragraph(
        doc,
        f"The UI workbook currently contains {len(ui_rows)} detailed UI test cases. These cases cover administrative and public-user flows, and every row includes a direct owner assignment, "
        "explicit preconditions, concise steps, expected result, actual result, final status, and evidence reference. The workbook therefore satisfies the course requirement for industry-standard testcase formatting.",
    )
    add_table(
        doc,
        "Table 3.5. Sample UI test case specification (positive case)",
        ["Field", "Value"],
        sample_case_rows(
            ui_case_pass,
            [
                "Test Case ID",
                "Scenario ID",
                "Title",
                "Module",
                "Owner",
                "Preconditions",
                "Test Data",
                "Steps",
                "Expected Result",
                "Actual Result",
                "Status",
                "Evidence / Note",
            ],
        ),
        [4.2, 11.8],
    )
    add_table(
        doc,
        "Table 3.6. Sample UI test case specification (defect-focused case)",
        ["Field", "Value"],
        sample_case_rows(
            ui_case_fail,
            [
                "Test Case ID",
                "Scenario ID",
                "Title",
                "Module",
                "Owner",
                "Preconditions",
                "Test Data",
                "Steps",
                "Expected Result",
                "Actual Result",
                "Status",
                "Evidence / Note",
            ],
        ),
        [4.2, 11.8],
    )
    add_figure(
        doc,
        img_invoice,
        "Figure 3.4. Defect-focused UI execution evidence for invoice creation failure.",
    )

    h = doc.add_paragraph(style="Heading 3")
    h.add_run("3.2.2. API test cases").bold = True
    add_paragraph(
        doc,
        f"The API workbook contains {len(api_rows)} API test cases derived from {scenario_interface_counts['API']} high-level API scenarios. "
        "The API cases use the same traceability principle as the UI cases, but add endpoint-specific fields such as method, authentication requirement, expected status code, and expected response body.",
    )
    add_table(
        doc,
        "Table 3.7. Sample API test case specification",
        ["Field", "Value"],
        sample_case_rows(
            api_case,
            [
                "Test Case ID",
                "Scenario ID",
                "Title",
                "Module",
                "Endpoint",
                "Method",
                "Auth Requirement",
                "Test Data",
                "Steps",
                "Expected Status Code",
                "Expected Body",
                "Actual Result",
                "Status",
                "Evidence / Note",
            ],
        ),
        [4.5, 11.5],
    )

    h = doc.add_paragraph(style="Heading 3")
    h.add_run("3.2.3. Test case allocation by member").bold = True
    add_paragraph(
        doc,
        "Ownership allocation was used to satisfy the requirement that each member contributes at least ten non-overlapping test cases. UI cases were distributed by module cluster, while API cases were executed as a shared QA-team artifact due to their smaller surface and automation-centric nature.",
    )
    add_table(
        doc,
        "Table 3.8. Test case allocation by member",
        ["Owner", "Assigned UI cases", "Allocation note"],
        [
            ["Hoang Van Thien", str(owner_counts["Hoang Van Thien"]), "Authentication, permissions, admin, invoices, reports, and partial public-catalog coverage"],
            ["Nguyen Thanh Dat", str(owner_counts["Nguyen Thanh Dat"]), "Customers, suppliers, purchases, and partial stock coverage"],
            ["Le Quang Duy", str(owner_counts["Le Quang Duy"]), "Products, product import, public catalog, and partial stock coverage"],
            ["QA Team", str(len(api_rows)), "Shared ownership for API execution and automation evidence"],
        ],
        [4.6, 3.2, 8.2],
    )
    add_paragraph(
        doc,
        "This allocation keeps module ownership visible while preserving a shared quality baseline across the final package. It also supports defense readiness because each member can explain both the rationale and the execution evidence of the areas they owned.",
    )

    h = doc.add_paragraph(style="Heading 3")
    h.add_run("3.2.4. Execution evidence summary").bold = True
    add_paragraph(
        doc,
        "The execution baseline is fully recorded in the final-results workbook and linked evidence files. Every designed test case in the current baseline has an execution result, and each fail case is traceable to evidence and defect records.",
    )
    add_table(
        doc,
        "Table 3.9. Execution evidence summary",
        ["Metric", "Value"],
        [
            ["Total executed test cases", str(len(final_rows))],
            ["Pass", str(final_exec_counts["Pass"])],
            ["Fail", str(final_exec_counts["Fail"])],
            ["Admin UI cases", str(final_interface_counts["Admin UI"])],
            ["Public UI cases", str(final_interface_counts["Public UI"])],
            ["API cases", str(final_interface_counts["API"])],
            ["Execution completeness", "100% of the current baseline"],
        ],
        [5.5, 10.5],
    )
    add_table(
        doc,
        "Table 3.10. Execution evidence types",
        ["Evidence type", "Purpose", "Representative artifact"],
        [
            ["UI screenshots", "Visual proof of page state and result outcome", "Admin login, import preview, invoice failure screenshots"],
            ["TRX runner output", "Execution trace for focused UI automation reruns", "UI rerun files under TestResults/RunnerOutput/UI"],
            ["Newman TXT and XML", "Exportable API automation proof", "newman-full-run.txt and newman-results.xml"],
            ["GitHub Issues", "Professional bug-management evidence", "Live issues for the four confirmed defects"],
            ["Automation video", "Presentation-ready execution walkthrough", "OSMS-Automation-Demo.mp4"],
        ],
        [4.0, 5.8, 6.2],
    )

    doc.save(out_path)
    print(out_path)
    print(out_path.stat().st_size)


if __name__ == "__main__":
    main()
