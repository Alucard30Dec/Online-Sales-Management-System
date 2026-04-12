from __future__ import annotations

import os
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt


def shade_cell(cell, fill: str = "D9EAF7") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def configure_document(doc: Document) -> None:
    sec = doc.sections[0]
    sec.top_margin = Cm(2.54)
    sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(2.54)
    sec.right_margin = Cm(2.54)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(6)

    for level, size in [(1, 16), (2, 14), (3, 13)]:
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = True
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(4)


def add_paragraph(doc: Document, text: str, align=None) -> None:
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="Normal")
        p.paragraph_format.left_indent = Cm(0.63)
        p.paragraph_format.first_line_indent = Cm(-0.3)
        run = p.add_run("• " + item)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def add_table(
    doc: Document,
    title: str,
    headers: list[str],
    rows: list[list[str]],
    col_widths: list[float] | None = None,
) -> None:
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(title)
    r.italic = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(11)
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        shade_cell(hdr[i])
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)

    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)

    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(11)
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.15
    doc.add_paragraph()


def add_figure(doc: Document, img_path: Path, caption: str, width_inches: float = 5.8) -> None:
    if not img_path.exists():
        return
    doc.add_picture(str(img_path), width=Inches(width_inches))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    r.italic = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(11)
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    doc.add_paragraph()


def main() -> None:
    out_dir = Path(os.environ["TARGET_DIR"])
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / "Chapter 2 - Test Plan.docx"

    img_modules = Path(os.environ["IMG_MODULES"])
    img_api = Path(os.environ["IMG_API"])
    img_import = Path(os.environ["IMG_IMPORT"])

    doc = Document()
    configure_document(doc)

    h = doc.add_paragraph(style="Heading 1")
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h.add_run("CHAPTER 2. TEST PLAN").bold = True

    h = doc.add_paragraph(style="Heading 2")
    h.add_run("2.1. Test Scope").bold = True
    add_paragraph(
        doc,
        "This chapter defines the boundaries of the testing effort for the Online Sales Management System. "
        "The scope was determined based on business impact, functional importance, visible user workflows, and "
        "source-informed risk analysis. Priority was given to areas where failures could directly affect access "
        "control, transaction correctness, inventory consistency, bulk data handling, and managerial visibility.",
    )
    add_paragraph(
        doc,
        "The final scope includes both user-facing and backend-accessible components. UI testing was prioritized "
        "for the administrative interface and the public product catalog, while API testing was used for the exposed "
        "service-health and catalog endpoints. The scope was designed to support structured manual execution, selective "
        "automation, and strong traceability between scenarios, test cases, results, evidence, and defect records.",
    )

    h = doc.add_paragraph(style="Heading 3")
    h.add_run("2.1.1. In-Scope").bold = True
    add_paragraph(doc, "The following items were included in the official testing scope of this submission:")
    add_bullets(
        doc,
        [
            "authentication and login validation",
            "role-based access control and permission boundaries",
            "admin user and admin group management",
            "customer and supplier master-data workflows",
            "product creation, update validation, duplicate checking, and product-state management",
            "product import validation, preview behavior, and import confirmation flow",
            "purchase creation, receiving, cancellation, and related validation behavior",
            "invoice creation, payment, cancellation, and inventory-related business behavior",
            "stock visibility, low-stock views, and stock-movement filtering",
            "report filtering and export-related source verification",
            "public product catalog search, filtering, sorting, and detail pages",
            "API verification for health and public catalog endpoints",
            "execution evidence collection, defect logging, metrics generation, and traceability mapping",
        ],
    )
    add_table(
        doc,
        "Table 2.1. In-scope test areas",
        ["Scope item", "Description"],
        [
            ["Admin UI", "Authentication, permissions, master data, purchases, invoices, stock, reports"],
            ["Public UI", "Product search, filtering, sorting, and detail pages"],
            ["API surface", "Health endpoint and public catalog endpoints"],
            ["Business validation", "Input validation, duplicate handling, role restrictions, inventory-sensitive rules"],
            ["Defect workflow", "Execution evidence, GitHub Issues, defect log, metrics, and result traceability"],
            ["Automation support", "UI and API automation for selected high-value flows"],
        ],
        [4.8, 11.2],
    )

    h = doc.add_paragraph(style="Heading 3")
    h.add_run("2.1.2. Out-of-Scope").bold = True
    add_paragraph(
        doc,
        "The following items were excluded from the scope of this submission due to time, environment, or project-boundary limitations:",
    )
    add_bullets(
        doc,
        [
            "performance and load testing",
            "penetration testing and deep security assessment",
            "mobile application testing",
            "full responsive certification across multiple mobile and tablet devices",
            "accessibility compliance audit",
            "disaster recovery, backup, and failover testing",
            "browser compatibility beyond the verified baseline and basic smoke validation",
            "undocumented write APIs outside the exposed catalog and health endpoints",
            "production deployment validation and cloud-infrastructure testing",
        ],
    )
    add_paragraph(
        doc,
        "These exclusions do not reduce the validity of the current testing scope. Instead, they clarify that the submission "
        "is focused on functional correctness, execution evidence, and business-rule verification rather than non-functional certification.",
    )
    add_table(
        doc,
        "Table 2.2. Out-of-scope items",
        ["Excluded item", "Reason for exclusion"],
        [
            ["Performance / load testing", "Not required by the course scope and no dedicated load environment was provided"],
            ["Deep security testing", "Requires specialized tools and a different testing objective"],
            ["Mobile / device matrix", "Not part of the required deliverables for this final submission"],
            ["Full cross-browser certification", "Only basic browser verification was feasible within the available environment"],
            ["Infrastructure / deployment testing", "Project was executed in a local testing environment, not a production-like deployment target"],
        ],
        [5.4, 10.6],
    )

    h = doc.add_paragraph(style="Heading 2")
    h.add_run("2.2. Test Strategy & Approach").bold = True
    add_paragraph(
        doc,
        "The testing strategy for this project was designed to balance breadth, realism, and evidence quality. Because the system "
        "includes both interactive business workflows and a small API surface, the team applied a mixed strategy consisting of manual execution, "
        "API execution, and basic automation. This combination made it possible to cover visible user behavior, backend-accessible responses, and "
        "repeatable high-value flows without overstating the maturity of the automation layer.",
    )
    add_paragraph(
        doc,
        "The strategy was primarily risk-based. Business-critical workflows such as authentication, permission boundaries, purchases, invoices, "
        "stock-sensitive operations, and product import were prioritized because defects in these areas could lead to data inconsistency, incorrect "
        "transaction behavior, or operational failures. Lower-risk areas such as cosmetic formatting and deep non-functional properties were not "
        "prioritized in the same way.",
    )

    h = doc.add_paragraph(style="Heading 3")
    h.add_run("2.2.1. Manual testing approach").bold = True
    add_paragraph(
        doc,
        "Manual testing was the primary execution method for the submission because it provides the broadest coverage for visible business behavior, "
        "validation feedback, access-control behavior, workflow branching, and exception handling. It was especially important for modules where human "
        "observation is necessary to verify form behavior, redirects, inline validation, data persistence, and functional correctness across multiple screens.",
    )
    add_paragraph(doc, "The manual approach followed a structured sequence:")
    add_bullets(
        doc,
        [
            "prepare the environment and seeded test data",
            "identify the target test case and required preconditions",
            "execute the defined steps exactly as documented",
            "compare actual system behavior against the expected result",
            "capture screenshots or output artifacts",
            "record the final result in the workbook",
            "log a defect when the mismatch is reproducible and supported by evidence",
        ],
    )

    h = doc.add_paragraph(style="Heading 3")
    h.add_run("2.2.2. Basic automation approach").bold = True
    add_paragraph(
        doc,
        "Basic automation was implemented to strengthen repeatability and bonus coverage rather than to replace the full manual test effort. "
        "Two automation layers were used in this project.",
    )
    add_paragraph(
        doc,
        "First, UI automation was implemented using Selenium WebDriver with .NET 8 and xUnit. This layer was used for selected high-value workflows "
        "such as login verification, permission checks, purchase creation, invoice-related validation, import preview behavior, and focused reruns "
        "for regression confirmation.",
    )
    add_paragraph(
        doc,
        "Second, API automation was implemented using Postman collections executed through Newman. This layer was used to verify health and catalog "
        "endpoints in a repeatable way and to produce exportable evidence such as text summaries and XML result files.",
    )
    add_paragraph(
        doc,
        "The automation scope was intentionally selective. The goal was to show practical, credible automation support with real outputs and evidence files, "
        "while avoiding exaggerated claims about full regression automation.",
    )
    add_table(
        doc,
        "Table 2.3. Automation strategy overview",
        ["Automation type", "Technology", "Main purpose"],
        [
            ["UI automation", "Selenium WebDriver + .NET 8 + xUnit", "Repeatable verification of selected high-value UI workflows"],
            ["API automation", "Postman + Newman", "Repeatable verification of exposed API endpoints and exportable runner evidence"],
            ["Evidence output", "Screenshots, TRX, TXT, XML", "Support result comparison, execution traceability, and bonus coverage"],
        ],
        [4.2, 5.8, 6.0],
    )
    add_figure(
        doc,
        img_api,
        "Figure 2.1. Representative API automation execution output using Newman.",
    )

    h = doc.add_paragraph(style="Heading 3")
    h.add_run("2.2.3. Black-box testing").bold = True
    add_paragraph(
        doc,
        "Black-box testing was the dominant testing strategy used in this project. Test execution focused on system inputs, user actions, visible outputs, "
        "and business outcomes without depending on internal implementation details during runtime verification.",
    )
    add_paragraph(doc, "This strategy was applied to:")
    add_bullets(
        doc,
        [
            "login success and login rejection",
            "role-based access restriction",
            "CRUD validation for products, suppliers, and customers",
            "purchase and invoice workflow outcomes",
            "product import validation and confirmation behavior",
            "public catalog search, sorting, and filtering behavior",
            "API status codes and payload validation",
        ],
    )
    add_paragraph(
        doc,
        "Black-box testing was suitable for the final submission because the course emphasizes practical test design, execution evidence, defect identification, "
        "and professional reporting. It also aligns well with user-visible functionality and real business requirements.",
    )

    h = doc.add_paragraph(style="Heading 3")
    h.add_run("2.2.4. White-box testing").bold = True
    add_paragraph(
        doc,
        "White-box testing was not used as a standalone implementation-heavy testing stream such as unit-test-driven structural coverage. However, "
        "white-box-informed analysis was used to support test design. The team inspected the source code, routes, controllers, validation logic, and seeded data "
        "in order to identify hidden edge cases, role conditions, unsupported inputs, and risky business-rule paths before executing the test cases.",
    )
    add_paragraph(doc, "This code-informed analysis helped the team design stronger scenarios for:")
    add_bullets(
        doc,
        [
            "permission boundaries across admin, sales, and warehouse roles",
            "invoice creation and cancellation behavior",
            "import validation rules and preview-confirm flow",
            "unsupported API query values",
            "route access and redirection logic",
            "stock-sensitive workflows and business-state transitions",
        ],
    )
    add_paragraph(
        doc,
        "Therefore, the project used a black-box execution model supported by limited white-box-informed test design. This is a practical and academically defensible "
        "use of both approaches for the available project scope.",
    )
    add_table(
        doc,
        "Table 2.4. Black-box and white-box usage",
        ["Approach", "How it was used in this project"],
        [
            ["Black-box", "Main execution strategy for UI and API verification based on inputs, actions, and visible outputs"],
            ["White-box-informed", "Used during test design through source inspection to identify edge cases, hidden branches, and risky logic"],
            ["Reason for combination", "Improves design quality without overstating the presence of deep structural automation or unit-level coverage"],
        ],
        [4.2, 11.8],
    )

    h = doc.add_paragraph(style="Heading 3")
    h.add_run("2.2.5. Entry and exit criteria").bold = True
    add_paragraph(
        doc,
        "Entry and exit criteria were defined to keep execution disciplined and to ensure that reported results were supported by reproducible evidence.",
    )
    add_paragraph(doc, "Entry criteria for execution:")
    add_bullets(
        doc,
        [
            "the application builds and runs successfully in the local environment",
            "the target base URL is reachable",
            "seeded accounts and required test data are available",
            "the required test case, preconditions, and test data are prepared",
            "the evidence folders and result workbooks are ready for recording outputs",
            "the target module can be opened in the environment being used",
        ],
    )
    add_paragraph(doc, "Exit criteria for the testing cycle:")
    add_bullets(
        doc,
        [
            "all designed UI and API test cases in the current baseline have been executed",
            "actual results and final statuses have been recorded",
            "screenshots, runner outputs, and evidence files have been stored",
            "confirmed mismatches have been logged into the defect register and GitHub Issues",
            "final results, metrics, and traceability files have been synchronized",
            "the report and presentation reflect the latest execution baseline",
        ],
    )
    add_paragraph(
        doc,
        "The execution cycle satisfied the defined exit criteria for coverage and documentation. However, product-quality closure was not fully achieved because "
        "four confirmed defects remained open at the final reporting point.",
    )
    add_table(
        doc,
        "Table 2.5. Entry and exit criteria",
        ["Criteria type", "Definition"],
        [
            ["Entry criteria", "Application is runnable, test data is available, target module is reachable, and evidence recording is ready"],
            ["Exit criteria", "All test cases are executed, results are recorded, defects are logged, and artifacts are synchronized"],
            ["Current status", "Execution exit criteria satisfied; product quality still limited by unresolved confirmed defects"],
        ],
        [4.5, 11.5],
    )

    h = doc.add_paragraph(style="Heading 2")
    h.add_run("2.3. Test Environment").bold = True
    add_paragraph(
        doc,
        "A controlled local test environment was used for all verified execution in this submission. The environment was configured to support both manual and "
        "automated testing, with the same application instance serving as the source for UI and API verification. Using a single consistent environment reduced "
        "result inconsistency and strengthened traceability across screenshots, runner outputs, workbooks, and defect records.",
    )
    add_paragraph(
        doc,
        "The environment section is important because it defines the reproducibility conditions of the submission. All reported outcomes in the final workbooks "
        "and screenshots were produced from this testing baseline unless explicitly stated otherwise.",
    )

    h = doc.add_paragraph(style="Heading 3")
    h.add_run("2.3.1. Hardware and software environment").bold = True
    add_paragraph(
        doc,
        "The execution environment used for the project was a Windows-based local workstation capable of running the ASP.NET Core MVC application, browser-based UI testing, "
        "Newman API execution, and supporting evidence tools. The application was executed locally using the .NET 8 runtime and connected to the database environment identified "
        "in the system as TiDB / test.",
    )
    add_paragraph(
        doc,
        "This environment was sufficient for the scope of the final submission because it supported functional execution, screenshot capture, UI automation, API execution, "
        "and evidence export. No distributed infrastructure or staging server was required for the tested endpoints and business flows included in the report.",
    )
    add_table(
        doc,
        "Table 2.6. Hardware and software environment",
        ["Environment item", "Value"],
        [
            ["Operating system", "Windows 11"],
            ["Application framework", "ASP.NET Core MVC (.NET 8)"],
            ["ORM / data access", "EF Core"],
            ["Authentication framework", "ASP.NET Identity"],
            ["Database tag observed in execution", "TiDB / test"],
            ["Execution type", "Local environment"],
            ["Base URL", "http://localhost:5068"],
            ["Repository path", "E:\\Project\\Online-Sales-Management-System"],
        ],
        [6.0, 10.0],
    )

    h = doc.add_paragraph(style="Heading 3")
    h.add_run("2.3.2. Browsers, OS, devices").bold = True
    add_paragraph(
        doc,
        "The primary browser used for UI execution was Google Chrome. Chrome was used for the main body of manual execution and for the majority of Selenium-based automated runs. "
        "In addition, a basic smoke rerun was executed on Microsoft Edge to provide minimal cross-browser evidence for the final submission.",
    )
    add_paragraph(
        doc,
        "Because the project is a web system and the course does not require a full mobile compatibility matrix, testing was performed on desktop conditions only. No mobile-device "
        "certification or tablet matrix was included. The browser strategy was therefore fit for purpose: one primary browser for complete functional verification and one secondary browser "
        "for basic compatibility evidence.",
    )
    add_table(
        doc,
        "Table 2.7. Browsers, OS, and devices",
        ["Item", "Usage in testing"],
        [
            ["Windows 11", "Main execution operating system"],
            ["Google Chrome", "Primary browser for manual UI and Selenium-based execution"],
            ["Microsoft Edge", "Basic smoke verification for limited cross-browser evidence"],
            ["Desktop / laptop environment", "Primary device class used for the final submission"],
            ["Mobile devices", "Not included in the official scope"],
            ["Tablet devices", "Not included in the official scope"],
        ],
        [5.2, 10.8],
    )
    add_figure(
        doc,
        img_modules,
        "Figure 2.2. Representative secured administrative environment and module access used during UI testing.",
    )

    h = doc.add_paragraph(style="Heading 3")
    h.add_run("2.3.3. Test accounts and test data").bold = True
    add_paragraph(
        doc,
        "The submission relied on seeded demo accounts and controlled test data rather than live production accounts. This approach reduced risk and made execution reproducible. "
        "The seeded accounts supported role-based verification across administrative, sales, and warehouse perspectives.",
    )
    add_paragraph(
        doc,
        "The test data set also included prepared values for product creation, purchase and invoice scenarios, API query combinations, and product import validation files. "
        "These datasets allowed the team to verify both happy-path and negative-path behavior in a controlled way.",
    )
    add_paragraph(
        doc,
        "Sensitive real-world credentials were not stored in the report. Only test-safe and seeded execution data was used for the final submission.",
    )
    add_table(
        doc,
        "Table 2.8. Test accounts and test data",
        ["Data category", "Example / source", "Purpose"],
        [
            ["Admin account", "admin@osms.local", "Authentication, admin access, invoice, report, and configuration-related workflows"],
            ["Sales account", "sales@osms.local", "Permission verification and sales-side behavior"],
            ["Warehouse account", "warehouse@osms.local", "Purchase and stock-related verification"],
            ["Negative login data", "Invalid password, unknown email, inactive-account scenarios", "Authentication validation"],
            ["Import dataset", "Mixed-validation workbook and invalid-format sample", "Product import preview and confirmation testing"],
            ["API query data", "Real category IDs, brand IDs, product IDs, and invalid parameters", "Catalog and health API verification"],
            ["Seeded business data", "Products, suppliers, customers, purchases, invoices, stock movement", "Business-flow execution support"],
        ],
        [4.2, 5.6, 6.2],
    )
    add_figure(
        doc,
        img_import,
        "Figure 2.3. Representative controlled validation dataset behavior during product import testing.",
    )

    h = doc.add_paragraph(style="Heading 3")
    h.add_run("2.3.4. Tools used").bold = True
    add_paragraph(
        doc,
        "The project used a practical set of tools chosen to match the requirements of the course and the structure of the system under test. No unnecessary tooling was introduced. "
        "Each tool supported a specific stage of the workflow, including planning, execution, automation, evidence collection, defect logging, and final reporting.",
    )
    add_paragraph(
        doc,
        "Excel-compatible files were used for traceable test case, result, and metrics management. Selenium WebDriver with xUnit was used for basic UI automation. "
        "Postman and Newman were used for API execution and exportable runner outputs. GitHub Issues was used as the real bug-management tool required by the course. "
        "Screenshots, TRX files, XML outputs, TXT logs, and the recorded automation video were used to support evidence completeness.",
    )
    add_table(
        doc,
        "Table 2.9. Tools used",
        ["Tool / technology", "Purpose"],
        [
            ["Microsoft Word", "Final report authoring and formatting"],
            ["Microsoft PowerPoint", "Final presentation preparation"],
            ["Excel-compatible XLSX / CSV", "Test cases, final results, metrics, and traceability"],
            ["Google Chrome", "Primary browser for UI execution"],
            ["Microsoft Edge", "Basic cross-browser smoke evidence"],
            ["Selenium WebDriver + .NET 8 + xUnit", "Basic UI automation"],
            ["Postman", "API request design and execution support"],
            ["Newman", "Automated API collection execution and exportable outputs"],
            ["GitHub Issues", "Defect tracking with severity, priority, and labels"],
            ["Screenshots / TRX / TXT / XML / MP4", "Execution evidence and automation proof"],
        ],
        [6.0, 10.0],
    )

    doc.save(out_path)
    print(out_path)
    print(out_path.stat().st_size)


if __name__ == "__main__":
    main()
