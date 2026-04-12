from __future__ import annotations

import csv
import os
from collections import Counter
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt


def shade_cell(cell, fill: str = "D9EAF7") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def configure_document(doc: Document) -> None:
    sec = doc.sections[0]
    sec.top_margin = Cm(2.54)
    sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(2.54)
    sec.right_margin = Cm(2.54)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(6)

    for level, size in [(1, 16), (2, 14), (3, 13)]:
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = True
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(4)


def add_paragraph(doc: Document, text: str, align=None) -> None:
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="Normal")
        p.paragraph_format.left_indent = Cm(0.63)
        p.paragraph_format.first_line_indent = Cm(-0.3)
        run = p.add_run("• " + item)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def add_table(
    doc: Document,
    title: str,
    headers: list[str],
    rows: list[list[str]],
    col_widths: list[float] | None = None,
) -> None:
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(title)
    r.italic = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(11)
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        shade_cell(hdr[i])
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)

    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)

    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(11)
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.15
    doc.add_paragraph()


def add_figure(doc: Document, img_path: Path, caption: str, width_inches: float = 5.8) -> None:
    if not img_path.exists():
        return
    doc.add_picture(str(img_path), width=Inches(width_inches))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    r.italic = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(11)
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    doc.add_paragraph()


def read_csv(path: Path) -> list[dict[str, str]]:
    with path.open(encoding="utf-8-sig", newline="") as f:
        return list(csv.DictReader(f))


def main() -> None:
    out_dir = Path(os.environ["TARGET_DIR"])
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / "Chapter 4 - Defect Report and Metrics.docx"

    defect_csv = Path(os.environ["DEFECT_CSV"])
    final_csv = Path(os.environ["FINAL_CSV"])
    metrics_csv = Path(os.environ["METRICS_CSV"])

    img_issue1 = Path(os.environ["IMG_ISSUE1"])
    img_issue2 = Path(os.environ["IMG_ISSUE2"])
    img_issue3 = Path(os.environ["IMG_ISSUE3"])
    img_issue4 = Path(os.environ["IMG_ISSUE4"])
    img_metrics = Path(os.environ["IMG_METRICS"])
    img_invoice_fail = Path(os.environ["IMG_INVOICE_FAIL"])

    defect_rows = read_csv(defect_csv)
    final_rows = read_csv(final_csv)
    metrics_rows = read_csv(metrics_csv)

    confirmed = [r for r in defect_rows if r["Record Type"] == "Confirmed Defect"]
    observations = [r for r in defect_rows if r["Record Type"] == "Observation"]
    high_critical = [r for r in confirmed if r["Severity"] in {"Critical", "High"}]
    medium_low = [r for r in confirmed if r["Severity"] in {"Medium", "Low"}]

    sev_counter = Counter(r["Severity"] for r in confirmed)
    prio_counter = Counter(r["Priority"] for r in confirmed)
    status_counter = Counter(r["Current Status"] for r in defect_rows)
    exec_counter = Counter(r["Execution Status"] for r in final_rows)

    total_exec = len(final_rows)
    fail_cases = exec_counter["Fail"]
    pass_cases = exec_counter["Pass"]
    blocked_cases = exec_counter.get("Blocked", 0)
    confirmed_count = len(confirmed)
    defect_ratio = round((confirmed_count / total_exec) * 100, 2) if total_exec else 0
    fail_rate = round((fail_cases / total_exec) * 100, 2) if total_exec else 0

    metrics_map = {row["Metric"]: row["Value"] for row in metrics_rows}

    doc = Document()
    configure_document(doc)

    h = doc.add_paragraph(style="Heading 1")
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h.add_run("CHAPTER 4. DEFECT REPORT & METRICS").bold = True

    h = doc.add_paragraph(style="Heading 2")
    h.add_run("4.1. Bug Management Tool and Workflow").bold = True
    add_paragraph(
        doc,
        "This chapter explains how defects were controlled, documented, and summarized during the testing process. Because the course requires a real bug-management tool, "
        "the project used GitHub Issues as the official issue-tracking platform for confirmed defects. Defect records were synchronized between the issue tracker, the defect register, "
        "the final results workbook, and the supporting screenshots and runner outputs.",
    )

    h = doc.add_paragraph(style="Heading 3")
    h.add_run("4.1.1. Selected tool").bold = True
    add_paragraph(
        doc,
        "GitHub Issues was selected as the official bug-management tool for this submission. The tool is appropriate for the project because it supports issue IDs, labels, screenshots, comments, "
        "status tracking, and direct linking from the final result package. It also satisfies the course requirement to use a real issue-management platform rather than a manually typed bug list only.",
    )
    add_paragraph(doc, "The selected tool was used for the four confirmed defects that remained reproducible in the current execution baseline.")
    add_table(
        doc,
        "Table 4.1. Bug-management tool selection",
        ["Item", "Value"],
        [
            ["Selected tool", "GitHub Issues"],
            ["Repository", "Alucard30Dec/Online-Sales-Management-System"],
            ["Used for", "Confirmed product defects with reproducible evidence"],
            ["Linked artifacts", "Defect register, final results, screenshots, TRX/TXT/XML evidence"],
            ["Reason selected", "Real issue IDs, practical traceability, and direct repository integration"],
        ],
        [4.8, 11.2],
    )
    add_figure(
        doc,
        img_issue1,
        "Figure 4.1. GitHub Issue example used as the official bug-management record for a confirmed defect.",
    )

    h = doc.add_paragraph(style="Heading 3")
    h.add_run("4.1.2. Issue workflow").bold = True
    add_paragraph(
        doc,
        "The issue workflow used in this submission was designed to keep defect reporting disciplined and evidence-based. An issue was created only after the mismatch was reproducible and supported by execution artifacts.",
    )
    add_bullets(
        doc,
        [
            "execute or rerun the related test case",
            "compare actual behavior against the expected result",
            "collect screenshots and runner outputs",
            "classify the result as observation or confirmed defect",
            "open a GitHub Issue for reproducible confirmed defects only",
            "link the issue back to the defect register and final results workbook",
            "keep the issue open until a real post-fix retest can verify closure",
        ],
    )
    add_table(
        doc,
        "Table 4.2. Issue workflow",
        ["Step", "Action", "Expected output"],
        [
            ["1", "Execute or rerun the test case", "Observable system result"],
            ["2", "Collect screenshot and runner evidence", "Supporting proof files"],
            ["3", "Compare expected and actual behavior", "Pass / Fail decision"],
            ["4", "Classify as observation or confirmed defect", "Defect triage decision"],
            ["5", "Create GitHub Issue for confirmed defects", "Live issue record with labels"],
            ["6", "Sync register, results, and evidence mapping", "Full traceability across artifacts"],
            ["7", "Retest after fix when available", "Potential closure or continued open status"],
        ],
        [1.5, 6.5, 8.0],
    )

    h = doc.add_paragraph(style="Heading 3")
    h.add_run("4.1.3. Severity and priority rules").bold = True
    add_paragraph(
        doc,
        "Severity and priority were assigned according to business impact, reproducibility, and operational risk. The project used simple but defensible rules so that each rating could be justified during presentation defense.",
    )
    add_table(
        doc,
        "Table 4.3. Severity and priority rules",
        ["Rule area", "How it was applied in this project"],
        [
            ["Severity: High", "Used when the defect blocks a core business workflow such as invoice creation, invoice cancellation, or valid product import confirmation"],
            ["Severity: Medium", "Used when the system blocks invalid input correctly but still presents broken or unreadable validation feedback"],
            ["Priority: High", "Used when the defect affects transaction correctness, inventory, or core operational workflow"],
            ["Priority: Medium", "Used when the defect should be fixed but does not fully break the core transaction path"],
        ],
        [4.0, 12.0],
    )
    add_paragraph(doc, "The actual issue labels used in GitHub include:")
    add_bullets(
        doc,
        [
            "severity:high / severity:medium",
            "priority:high / priority:medium",
            "status:open",
            "module:invoices / module:purchases / module:product-import",
            "interface:web-ui",
            "type:defect",
        ],
    )

    h = doc.add_paragraph(style="Heading 2")
    h.add_run("4.2. Defect Log").bold = True
    add_paragraph(
        doc,
        f"At the current reporting point, the defect register contains {len(confirmed)} confirmed defects and {len(observations)} closed observations. "
        "The observations are retained for historical traceability, but only the confirmed defects are counted as current product defects in the summary metrics.",
    )

    h = doc.add_paragraph(style="Heading 3")
    h.add_run("4.2.1. Critical and high severity defects").bold = True
    add_paragraph(
        doc,
        "No active critical-severity defect exists in the current baseline. However, three high-severity defects remain open and affect core business workflows. "
        "These defects are the most important findings of the submission because they directly affect invoice processing, inventory-related behavior, or valid import completion.",
    )
    add_table(
        doc,
        "Table 4.4. High-severity confirmed defects",
        ["Defect ID", "Related test case(s)", "Module", "Severity", "Priority", "Current status", "GitHub Issue"],
        [
            [
                r["Record ID"],
                r["Related Test Case ID"],
                r["Module"],
                r["Severity"],
                r["Priority"],
                r["Current Status"],
                r["GitHub Issue"],
            ]
            for r in high_critical
        ],
        [2.8, 4.4, 2.6, 1.6, 1.6, 2.6, 5.0],
    )
    add_figure(
        doc,
        img_invoice_fail,
        "Figure 4.2. Reproduction evidence for the high-severity invoice creation defect.",
    )
    add_figure(
        doc,
        img_issue4,
        "Figure 4.3. GitHub Issue evidence for a high-severity invoice cancellation defect.",
    )

    h = doc.add_paragraph(style="Heading 3")
    h.add_run("4.2.2. Medium and low severity defects").bold = True
    add_paragraph(
        doc,
        "The current baseline contains one medium-severity confirmed defect and no low-severity confirmed defect. The medium defect does not break transaction blocking logic itself, "
        "but it still degrades the quality of system feedback by rendering an unreadable validation banner in the purchase-create flow.",
    )
    add_table(
        doc,
        "Table 4.5. Medium and low severity confirmed defects",
        ["Defect ID", "Related test case(s)", "Module", "Severity", "Priority", "Current status", "GitHub Issue"],
        [
            [
                r["Record ID"],
                r["Related Test Case ID"],
                r["Module"],
                r["Severity"],
                r["Priority"],
                r["Current Status"],
                r["GitHub Issue"],
            ]
            for r in medium_low
        ]
        + [["Low severity defects", "None", "-", "Low", "-", "-", "-"]],
        [2.8, 4.4, 2.6, 1.6, 1.6, 2.6, 5.0],
    )
    add_figure(
        doc,
        img_issue2,
        "Figure 4.4. GitHub Issue evidence for the medium-severity purchase validation defect.",
    )

    h = doc.add_paragraph(style="Heading 3")
    h.add_run("4.2.3. Screenshots and reproduction evidence").bold = True
    add_paragraph(
        doc,
        "Every confirmed defect in the current baseline is supported by reproduction steps and at least one screenshot or runner artifact. The strongest defect packages also include retest evidence and direct GitHub Issue screenshots. "
        "This evidence design improves traceability and makes the defect records defensible during presentation review.",
    )
    add_table(
        doc,
        "Table 4.6. Defect evidence mapping summary",
        ["Defect ID", "Evidence types used", "Purpose"],
        [
            [
                "BUG-20260406-001",
                "UI screenshots, TRX, server log excerpt, GitHub Issue screenshot",
                "Proves invoice-create failure is reproducible and not limited to one scenario",
            ],
            [
                "BUG-20260411-002",
                "UI screenshots, focused rerun TRX, GitHub Issue screenshot",
                "Proves purchase validation banner defect is reproducible",
            ],
            [
                "BUG-20260411-003",
                "UI screenshots, focused rerun TRX, GitHub Issue screenshot",
                "Proves import confirmation still fails after valid preview",
            ],
            [
                "BUG-20260411-004",
                "UI screenshots, focused rerun TRX, GitHub Issue screenshot",
                "Proves invoice cancellation still fails and stock return does not occur",
            ],
        ],
        [2.8, 5.8, 7.4],
    )
    add_figure(
        doc,
        img_issue3,
        "Figure 4.5. GitHub Issue evidence for the product-import confirmed defect.",
    )

    h = doc.add_paragraph(style="Heading 2")
    h.add_run("4.3. Test Summary Metrics").bold = True
    add_paragraph(
        doc,
        "The test summary metrics describe the final execution state of the current baseline. It is important to distinguish between failing test cases and unique confirmed defects. "
        "In this project, 7 test cases currently report Fail, but those failures map to 4 unique confirmed defects because multiple failing cases are caused by the same underlying issue.",
    )

    h = doc.add_paragraph(style="Heading 3")
    h.add_run("4.3.1. Total test cases executed").bold = True
    add_paragraph(
        doc,
        f"The current baseline contains {total_exec} executed test cases. According to the synchronized final-results workbook, all designed UI and API cases in the baseline have recorded execution results, "
        f"which means execution completeness is {metrics_map.get('Execution Progress %', '100')}%.",
    )
    add_table(
        doc,
        "Table 4.7. Execution totals",
        ["Metric", "Value"],
        [
            ["Total test cases", str(total_exec)],
            ["Executed test cases", metrics_map.get("Executed Test Cases", str(total_exec))],
            ["Execution progress", metrics_map.get("Execution Progress %", "100") + "%"],
        ],
        [6.0, 10.0],
    )

    h = doc.add_paragraph(style="Heading 3")
    h.add_run("4.3.2. Pass/Fail/Blocked statistics").bold = True
    add_paragraph(
        doc,
        f"The current execution baseline contains {pass_cases} passing test cases, {fail_cases} failing test cases, and {blocked_cases} blocked test cases. "
        f"This gives a pass rate of {metrics_map.get('Pass Rate On Executed %', '88.89')}% and a fail rate of {metrics_map.get('Fail Rate On Executed %', '11.11')}%.",
    )
    add_table(
        doc,
        "Table 4.8. Pass / Fail / Blocked statistics",
        ["Status", "Count", "Interpretation"],
        [
            ["Pass", str(pass_cases), "Expected behavior matched real execution evidence"],
            ["Fail", str(fail_cases), "Observed behavior did not match the expected result"],
            ["Blocked", str(blocked_cases), "No blocked cases remain in the current baseline"],
        ],
        [3.2, 2.0, 10.8],
    )

    h = doc.add_paragraph(style="Heading 3")
    h.add_run("4.3.3. Defect count by severity").bold = True
    add_paragraph(
        doc,
        "Severity distribution should be interpreted using confirmed defects only, not raw failing test cases. The current confirmed-defect profile shows concentration in high-severity business workflows, "
        "especially invoice and product-import related behavior.",
    )
    add_table(
        doc,
        "Table 4.9. Confirmed defect count by severity",
        ["Severity", "Confirmed defect count"],
        [
            ["Critical", str(sev_counter.get("Critical", 0))],
            ["High", str(sev_counter.get("High", 0))],
            ["Medium", str(sev_counter.get("Medium", 0))],
            ["Low", str(sev_counter.get("Low", 0))],
        ],
        [6.0, 10.0],
    )

    h = doc.add_paragraph(style="Heading 3")
    h.add_run("4.3.4. Defect ratio and observations").bold = True
    add_paragraph(
        doc,
        f"The confirmed-defect ratio in the current baseline is {defect_ratio}%, calculated as {confirmed_count} confirmed defects out of {total_exec} executed test cases. "
        f"This ratio is lower than the fail-case rate of {fail_rate}% because several failing test cases map to the same underlying defect record. "
        "This distinction is important for accurate quality reporting.",
    )
    add_paragraph(
        doc,
        f"In addition to the confirmed defects, the defect register also records {len(observations)} observations that were closed after rerun analysis. "
        "These records remain useful because they show that the team did not automatically classify every unexpected behavior as a product defect. Instead, each mismatch was reviewed and reclassified based on reproducible evidence.",
    )
    add_table(
        doc,
        "Table 4.10. Defect ratio and observations",
        ["Metric", "Value", "Note"],
        [
            ["Confirmed defects", str(confirmed_count), "Unique product defects currently open"],
            ["Failing test cases", str(fail_cases), "Some failures share the same root cause"],
            ["Defect ratio", f"{defect_ratio}%", "Confirmed defects / executed test cases"],
            ["Fail-case rate", f"{fail_rate}%", "Failing cases / executed test cases"],
            ["Closed observations", str(len(observations)), "Historical mismatch records resolved by rerun analysis"],
        ],
        [4.5, 2.5, 9.0],
    )
    add_figure(
        doc,
        img_metrics,
        "Figure 4.6. Metrics summary exported from the synchronized results workbook.",
    )

    doc.save(out_path)
    print(out_path)
    print(out_path.stat().st_size)


if __name__ == "__main__":
    main()
