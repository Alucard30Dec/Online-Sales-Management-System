from __future__ import annotations

import csv
import os
from collections import Counter
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


def shade_cell(cell, fill: str = "D9EAF7") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def configure_document(doc: Document) -> None:
    sec = doc.sections[0]
    sec.top_margin = Cm(2.54)
    sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(2.54)
    sec.right_margin = Cm(2.54)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(6)

    for level, size in [(1, 16), (2, 14), (3, 13)]:
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = True
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(4)


def add_paragraph(doc: Document, text: str, align=None, bold: bool = False) -> None:
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run.bold = bold
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="Normal")
        p.paragraph_format.left_indent = Cm(0.63)
        p.paragraph_format.first_line_indent = Cm(-0.3)
        run = p.add_run("- " + item)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def add_table(
    doc: Document,
    title: str,
    headers: list[str],
    rows: list[list[str]],
    col_widths: list[float] | None = None,
) -> None:
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(title)
    r.italic = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(11)
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        shade_cell(hdr[i])
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)

    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)

    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(11)
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.15
    doc.add_paragraph()


def read_csv(path: Path) -> list[dict[str, str]]:
    with path.open(encoding="utf-8-sig", newline="") as f:
        return list(csv.DictReader(f))


def main() -> None:
    out_dir = Path(os.environ["TARGET_DIR"])
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / "Appendix - References - Task Division.docx"

    clean_root = Path(os.environ["CLEAN_ROOT"])
    ui_csv = Path(os.environ["UI_CSV"])
    api_csv = Path(os.environ["API_CSV"])
    final_csv = Path(os.environ["FINAL_CSV"])
    defect_csv = Path(os.environ["DEFECT_CSV"])

    ui_rows = read_csv(ui_csv)
    api_rows = read_csv(api_csv)
    final_rows = read_csv(final_csv)
    defect_rows = read_csv(defect_csv)

    owner_counts = Counter(r["Owner"] for r in ui_rows)
    total_ui = len(ui_rows)
    pass_fail = Counter(r["Execution Status"] for r in final_rows)
    confirmed_defects = [r for r in defect_rows if r["Record Type"] == "Confirmed Defect"]

    github_repo = "https://github.com/Alucard30Dec/Online-Sales-Management-System"
    github_clean = "https://github.com/Alucard30Dec/Online-Sales-Management-System/tree/main/Report%20Test%20subject/SV00123-ATU-A01"

    doc = Document()
    configure_document(doc)

    h = doc.add_paragraph(style="Heading 1")
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h.add_run("APPENDIX").bold = True

    h = doc.add_paragraph(style="Heading 2")
    h.add_run("Appendix 1. GitHub").bold = True
    add_paragraph(
        doc,
        "The final submission package is stored in the project repository under the clean submission folder so that the examiner can access the report, test cases, results, automation assets, and video from a single structured location.",
    )
    add_table(
        doc,
        "Table A.1. GitHub submission links",
        ["Item", "Link / location"],
        [
            ["GitHub Repository", github_repo],
            ["Clean submission package", github_clean],
            ["Top-level folders in clean package", ", ".join(sorted(p.name for p in clean_root.iterdir()))],
        ],
        [4.8, 11.2],
    )

    h = doc.add_paragraph(style="Heading 2")
    h.add_run("Appendix 2. Excel test cases").bold = True
    add_paragraph(
        doc,
        "The official test-case artifacts are maintained in Excel-compatible workbook and CSV form. These files are the source of truth for scenario-to-case mapping, test ownership, execution details, and traceability across the final submission package.",
    )
    add_table(
        doc,
        "Table A.2. Excel and scenario artifacts",
        ["Artifact", "Package-relative path", "Purpose"],
        [
            ["UI test cases workbook", "TestCases/UI/OSMS-UI-Test-Cases.xlsx", f"{len(ui_rows)} UI test cases with ownership, expected and actual results, and status"],
            ["API test cases workbook", "TestCases/API/OSMS-API-Test-Cases.xlsx", f"{len(api_rows)} API test cases with request and result coverage"],
            ["UI test cases CSV", "TestCases/UI/OSMS-UI-Test-Cases.csv", "CSV export for traceability and automation-friendly review"],
            ["API test cases CSV", "TestCases/API/OSMS-API-Test-Cases.csv", "CSV export for API case review"],
            ["Scenario list", "TestCases/Scenarios/test-scenarios.md", "Documented UI and API scenario set used to derive the test cases"],
        ],
        [4.3, 6.1, 5.6],
    )
    add_paragraph(
        doc,
        "Figma design link: No verified Figma URL was found in the current repository artifacts. If the team has a real shared Figma design link, it should be inserted manually here; otherwise this item should be omitted rather than claimed without evidence.",
    )

    h = doc.add_paragraph(style="Heading 2")
    h.add_run("Appendix 3. Test script and test data").bold = True
    add_paragraph(
        doc,
        "The automation scripts and prepared test data used in the final submission are stored in the TestScript-Data folder. These artifacts support manual reruns, UI automation, API automation, and data-driven validation scenarios such as product import.",
    )
    add_table(
        doc,
        "Table A.3. Test scripts and test data",
        ["Artifact group", "Package-relative path", "Notes"],
        [
            ["Automation root", "TestScript-Data/Automation/", "Contains UI and API automation assets"],
            ["UI automation", "TestScript-Data/Automation/ui/", "Selenium WebDriver + .NET 8 + xUnit test project"],
            ["API automation", "TestScript-Data/Automation/api/", "Postman collection, environment, and Newman run script"],
            ["Test data root", "TestScript-Data/TestData/", "Prepared seeded data and import files"],
            ["Account data", "TestScript-Data/TestData/accounts/OSMS-Test-Accounts.md", "Seeded accounts used for role-based verification"],
            ["Import data", "TestScript-Data/TestData/ui/", "Mixed-validation and negative import files"],
            ["API data", "TestScript-Data/TestData/api/OSMS-API-Test-Data.json", "Reference values for API execution"],
        ],
        [4.2, 6.1, 5.7],
    )
    add_paragraph(doc, "Verified tools used in the submission package:")
    add_bullets(
        doc,
        [
            "Visual Studio / .NET 8 tooling for application build and UI automation execution",
            "Google Chrome and Microsoft Edge for browser-based execution",
            "Postman and Newman for API execution",
            "GitHub Issues for defect management",
            "Excel-compatible workbooks and CSV files for test cases, final results, and metrics",
            "Microsoft Word and PowerPoint for the final report and presentation deliverables",
        ],
    )
    add_paragraph(doc, "Generative AI disclosure:", bold=True)
    add_bullets(
        doc,
        [
            "Used for: summarizing requirements, suggesting naming conventions, improving English grammar in the report, and providing code guidance under team review.",
            "Not used for: direct copy-paste of full modules or fabricated testing evidence without review.",
            "Verification: all submission-facing outputs were checked manually through team review, execution evidence, and workbook synchronization before inclusion.",
        ],
    )

    h = doc.add_paragraph(style="Heading 2")
    h.add_run("Appendix 4. Final results workbook and images").bold = True
    add_paragraph(
        doc,
        "The final execution baseline is documented in the result workbook and its supporting evidence folders. These files are used to justify the pass, fail, defect, and traceability claims made in the report.",
    )
    add_table(
        doc,
        "Table A.4. Final results and evidence package",
        ["Artifact", "Package-relative path", "Purpose"],
        [
            ["Final results workbook", "TestResults/FinalResults/OSMS-Final-Test-Results.xlsx", f"Records {len(final_rows)} executed cases with final status"],
            ["Final results CSV", "TestResults/FinalResults/OSMS-Final-Results.csv", f"Summary baseline: {pass_fail['Pass']} Pass / {pass_fail['Fail']} Fail"],
            ["Execution-evidence mapping", "TestResults/FinalResults/execution-evidence-mapping.csv", "Links test cases to screenshots, runner outputs, and issues"],
            ["Metrics workbook", "TestResults/Metrics/OSMS-Test-Metrics.xlsx", "Aggregated metrics and summary views"],
            ["UI evidence", "TestResults/Evidence/UI/automation/", "Screenshots proving UI execution outcomes"],
            ["API evidence", "TestResults/Evidence/API/newman-full-run.txt", "Text summary of Newman API execution"],
            ["Report evidence", "TestResults/Evidence/Report/", "Screenshots of metrics and report-support visuals"],
        ],
        [4.4, 6.0, 5.6],
    )

    h = doc.add_paragraph(style="Heading 2")
    h.add_run("Appendix 5. Automation video").bold = True
    add_paragraph(
        doc,
        "The submission includes a recorded automation video used to demonstrate UI and API execution evidence in a concise presentation-friendly format.",
    )
    add_table(
        doc,
        "Table A.5. Automation video artifacts",
        ["Artifact", "Package-relative path", "Purpose"],
        [
            ["Automation demo video", "Videos/OSMS-Automation-Demo.mp4", "Primary automation demonstration artifact"],
            ["Video description", "Videos/Automation-Video-Description.md", "Explains covered flows and evidence mapping"],
            ["Recording script", "Videos/record-automation-demo.ps1", "Repeatable capture workflow for the demo package"],
        ],
        [4.5, 5.8, 5.7],
    )

    h = doc.add_paragraph(style="Heading 2")
    h.add_run("Appendix 6. Bug tracker screenshots").bold = True
    add_paragraph(
        doc,
        "Bug-management evidence is stored as both workbook and screenshot artifacts. The defect package combines the defect register, the workbook log, and four live GitHub Issues for the confirmed defects in the current baseline.",
    )
    add_table(
        doc,
        "Table A.6. Bug tracker and defect artifacts",
        ["Artifact", "Package-relative path", "Notes"],
        [
            ["Defect register CSV", "TestResults/Defects/OSMS-Defect-Register.csv", f"{len(confirmed_defects)} confirmed defects plus closed observations"],
            ["Defect log workbook", "TestResults/Defects/OSMS-Defect-Log.xlsx", "Submission-facing defect tracking workbook"],
            ["GitHub issue screenshots", "TestResults/Evidence/Defects/", "Screenshots for issues #1, #2, #3, and #4"],
            ["GitHub issue checklist files", "TestResults/Defects/GitHubIssues/", "Structured issue drafts and evidence notes"],
        ],
        [4.7, 5.8, 5.5],
    )

    doc.add_page_break()

    h = doc.add_paragraph(style="Heading 1")
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h.add_run("REFERENCES").bold = True

    add_paragraph(
        doc,
        "Bootstrap. (n.d.). Get started with Bootstrap. Bootstrap v5.3. Retrieved April 12, 2026, from https://getbootstrap.com/docs/5.3/getting-started/introduction/",
    )
    add_paragraph(
        doc,
        "GitHub. (n.d.). About Git. GitHub Docs. Retrieved April 12, 2026, from https://docs.github.com/en/get-started/using-git/about-git",
    )
    add_paragraph(
        doc,
        "Microsoft. (2023, January 12). Migrations overview. Microsoft Learn. https://learn.microsoft.com/en-us/ef/core/managing-schemas/migrations/",
    )
    add_paragraph(
        doc,
        "Microsoft. (2024, June 17). Overview of ASP.NET Core MVC. Microsoft Learn. https://learn.microsoft.com/en-us/aspnet/core/mvc/overview",
    )
    add_paragraph(
        doc,
        "Microsoft. (2024, November 12). Overview of Entity Framework Core. Microsoft Learn. https://learn.microsoft.com/en-us/ef/core/",
    )
    add_paragraph(
        doc,
        "Microsoft. (2025, July 30). Overview of ASP.NET Core. Microsoft Learn. https://learn.microsoft.com/en-us/aspnet/core/overview",
    )
    add_paragraph(
        doc,
        "Microsoft. (2025, November 12). Introduction to Identity on ASP.NET Core. Microsoft Learn. https://learn.microsoft.com/en-us/aspnet/core/security/authentication/identity",
    )

    doc.add_page_break()

    h = doc.add_paragraph(style="Heading 1")
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h.add_run("TASK DIVISION AND CONTRIBUTION SUMMARY").bold = True

    add_paragraph(
        doc,
        "The following summary reflects the primary UI test-case ownership used in the final baseline, together with the main integration responsibilities observed in the submission package. "
        "The workload percentage is derived from the non-overlapping UI ownership counts because that is the most directly verifiable per-member distribution in the current artifacts.",
    )

    owner_rows = []
    role_map = {
        "Hoang Van Thien": "Project audit, traceability alignment, final integration, authentication, permissions, admin management, invoices, reports, partial public catalog",
        "Nguyen Thanh Dat": "Customers, suppliers, purchases, partial stock workflows",
        "Le Quang Duy": "Products, product import, partial stock, public catalog workflows",
    }
    display_map = {
        "Hoang Van Thien": "Hoàng Văn Thiên",
        "Nguyen Thanh Dat": "Nguyễn Thành Đạt",
        "Le Quang Duy": "Lê Quang Duy",
    }
    for owner in ["Hoang Van Thien", "Nguyen Thanh Dat", "Le Quang Duy"]:
        count = owner_counts.get(owner, 0)
        pct = round((count / total_ui) * 100, 2) if total_ui else 0
        owner_rows.append(
            [
                display_map[owner],
                role_map[owner],
                f"{pct:.2f}%",
                "100%",
            ]
        )

    add_table(
        doc,
        "Table T.1. Task division and contribution summary",
        ["Member name", "Assigned tasks", "Workload contribution (%)", "Completion status (%)"],
        owner_rows,
        [3.2, 8.4, 3.0, 3.0],
    )
    add_paragraph(
        doc,
        "Note: the workload contribution percentage above is calculated from primary UI case ownership only. Shared API execution, report consolidation, and packaging activities were performed collaboratively and are therefore described in the assigned-task column rather than re-weighted numerically.",
    )

    doc.save(out_path)
    print(out_path)
    print(out_path.stat().st_size)


if __name__ == "__main__":
    main()
